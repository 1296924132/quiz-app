#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
创建已填写的2026年暑期社会实践活动院级项目申报书（改进版）
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ========== 设置默认字体 ==========
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
pf.space_before = Pt(0)
pf.space_after = Pt(0)

# ========== 页面设置 ==========
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# ========== 辅助函数 ==========
def set_run_font(run, name='宋体', size=Pt(12), bold=False):
    """设置run的字体"""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_centered_para(text, font_name='宋体', size=Pt(12), bold=False, space_before=0, space_after=0):
    """添加居中段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(28)
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return p

def add_left_para(text, font_name='仿宋', size=Pt(12), bold=False, space_before=0, space_after=0, indent=0):
    """添加左对齐段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(24)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return p

def set_cell_text(cell, text, bold=False, size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格文本"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    set_run_font(run, '宋体', size, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_cell_left(cell, text, bold=False, size=Pt(10)):
    set_cell_text(cell, text, bold, size, WD_ALIGN_PARAGRAPH.LEFT)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" '
            f'w:color="{val.get("color", "000000")}" '
            f'w:space="0"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def merge_cells(table, row_start, col_start, row_end, col_end):
    """合并单元格"""
    cell = table.cell(row_start, col_start)
    for r in range(row_start, row_end + 1):
        for c in range(col_start, col_end + 1):
            if r != row_start or c != col_start:
                cell = cell.merge(table.cell(r, c))
    return cell

# ========== 封 面 ==========
for _ in range(5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)

# 主标题
add_centered_para('2026年暑期社会实践活动', '黑体', Pt(26), True, 0, 6)
add_centered_para('院 级 项 目 申 报 书', '黑体', Pt(26), True, 6, 0)

for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)

# 封面信息
cover_fields = ['团队名称：', '项目名称：', '所在单位：', '负责人姓名：']
for field in cover_fields:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(36)
    run = p.add_run(field)
    set_run_font(run, '仿宋', Pt(16))
    run2 = p.add_run('_______________')
    set_run_font(run2, '仿宋', Pt(16))

for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)

# 落款
add_centered_para('共青团哈尔滨工程大学委员会', '仿宋', Pt(16), False, 0, 6)
add_centered_para('2026年6月', '仿宋', Pt(16), False, 6, 0)

# ========== A 基本情况 ==========
doc.add_page_break()

add_left_para('A  基本情况', '黑体', Pt(14), True, 6, 6)

# ---- 主表格（团队情况+负责人+指导教师）----
main_table = doc.add_table(rows=5, cols=7)
main_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 设置列宽 (cm)
col_widths = [2.0, 2.2, 2.0, 2.0, 2.0, 2.2, 2.2]
for i, w in enumerate(col_widths):
    for cell in main_table.columns[i].cells:
        cell.width = Cm(w)

# === 第0行：团队情况 ===
merge_cells(main_table, 0, 1, 0, 6)  # 合并右侧
set_cell_text(main_table.rows[0].cells[0], '团 队\n情 况', True, Pt(9))
set_cell_left(main_table.rows[0].cells[1],
    '团队名称：例：哈尔滨工程大学________（学院）赴________省________（市）暑期社会实践团\n'
    '实践地点：__________________________________\n'
    '参与人数：______人    活动起止时间：2026年__月__日 — 2026年__月__日\n'
    '预算资金：________元    自筹或资助资金：________元\n'
    '项目全称：__________________________________',
    False, Pt(9))

# === 第1行：负责人 ===
merge_cells(main_table, 1, 1, 1, 6)
set_cell_text(main_table.rows[1].cells[0], '负 责 人', True, Pt(9))
set_cell_left(main_table.rows[1].cells[1],
    '姓  名：________    性  别：________    政治面貌：________\n'
    '学  院：________    专  业：________    学  号：________\n'
    '联系方式：________    E-mail：__________________',
    False, Pt(9))

# === 第2行：指导教师 ===
merge_cells(main_table, 2, 1, 2, 6)
set_cell_text(main_table.rows[2].cells[0], '指导\n教师', True, Pt(9))
set_cell_left(main_table.rows[2].cells[1],
    '姓  名：________    单  位：________    专  业：________\n'
    '职称/职务：________    联系方式：________\n'
    '本人签字：__________________',
    False, Pt(9))

# === 第3行：主要成员（表头） ===
# 注意：不要合并右侧列，保持各列独立以显示表头
set_cell_text(main_table.rows[3].cells[0], '主 要\n成 员', True, Pt(9))
set_cell_text(main_table.rows[3].cells[1], '姓名', True, Pt(9))
set_cell_text(main_table.rows[3].cells[2], '政治面貌', True, Pt(9))
set_cell_text(main_table.rows[3].cells[3], '学院', True, Pt(9))
set_cell_text(main_table.rows[3].cells[4], '专业', True, Pt(9))
set_cell_text(main_table.rows[3].cells[5], '学号', True, Pt(9))
set_cell_text(main_table.rows[3].cells[6], '本人签字', True, Pt(9))

# === 第4行：成员数据行 ===
set_cell_text(main_table.rows[4].cells[0], '', True, Pt(9))
for i in range(1, 7):
    set_cell_text(main_table.rows[4].cells[i], '', False, Pt(9))

# ========== B 实践计划 ==========
doc.add_page_break()

add_left_para('B  实践计划', '黑体', Pt(14), True, 6, 12)

add_left_para('一、活动简要计划', '黑体', Pt(12), True, 6, 6)
for i in range(6):
    add_left_para('（' + str(i+1) + '）____________________________________________________________', '仿宋', Pt(11), False, 2, 2)

add_left_para('二、活动内容', '黑体', Pt(12), True, 12, 6)
for i in range(4):
    add_left_para('___________________________________________________________________________', '仿宋', Pt(11), False, 2, 2)

add_left_para('三、活动组织方式构想', '黑体', Pt(12), True, 12, 6)
for i in range(4):
    add_left_para('___________________________________________________________________________', '仿宋', Pt(11), False, 2, 2)

add_left_para('四、经费预算', '黑体', Pt(12), True, 12, 6)

# 经费预算表
budget_table = doc.add_table(rows=9, cols=2)
budget_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in budget_table.columns[0].cells:
    cell.width = Cm(11)
for cell in budget_table.columns[1].cells:
    cell.width = Cm(4)

set_cell_text(budget_table.rows[0].cells[0], '用    途', True, Pt(10))
set_cell_text(budget_table.rows[0].cells[1], '经费小计（元）', True, Pt(10))

items = ['交通费', '住宿费', '餐饮补贴', '资料打印费', '保险费', '宣传费', '其他费用']
for i, item in enumerate(items):
    set_cell_left(budget_table.rows[i+1].cells[0], item, False, Pt(10))
    set_cell_text(budget_table.rows[i+1].cells[1], '', False, Pt(10))

set_cell_text(budget_table.rows[8].cells[0], '合  计', True, Pt(10))
set_cell_text(budget_table.rows[8].cells[1], '', True, Pt(10))

# ========== C 实践项目可行性分析 ==========
doc.add_page_break()

add_left_para('C  实践项目可行性分析', '黑体', Pt(14), True, 6, 12)

add_left_para('一、实践项目研究内容', '黑体', Pt(12), True, 6, 6)
for i in range(8):
    add_left_para('___________________________________________________________________________', '仿宋', Pt(11), False, 2, 2)

add_left_para('二、项目意义与目的', '黑体', Pt(12), True, 12, 6)
for i in range(6):
    add_left_para('___________________________________________________________________________', '仿宋', Pt(11), False, 2, 2)

add_left_para('三、项目研究的基本思路', '黑体', Pt(12), True, 12, 6)
for i in range(6):
    add_left_para('___________________________________________________________________________', '仿宋', Pt(11), False, 2, 2)

# ========== D 项目审核意见 ==========
doc.add_page_break()

add_left_para('D  项目审核意见', '黑体', Pt(14), True, 6, 12)

# 审核意见表
review_table = doc.add_table(rows=3, cols=2)
review_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in review_table.columns[0].cells:
    cell.width = Cm(2.5)
for cell in review_table.columns[1].cells:
    cell.width = Cm(12.5)

# 第0行：单位意见标题
set_cell_text(review_table.rows[0].cells[0], '单 位\n意 见', True, Pt(10))
set_cell_left(review_table.rows[0].cells[1],
    '非团队负责人所在学院团队成员外出实践审核意见：\n\n'
    '姓  名            学  院            学  号            单位意见（签字）\n'
    '________        ________        ________        ________________\n'
    '________        ________        ________        ________________\n'
    '________        ________        ________        ________________',
    False, Pt(10))

# 第1行：空白行（用作格式）
set_cell_text(review_table.rows[1].cells[0], '', False, Pt(10))
set_cell_left(review_table.rows[1].cells[1], '', False, Pt(10))

# 第2行：团队负责人所在单位意见
set_cell_text(review_table.rows[2].cells[0], '', False, Pt(10))
set_cell_left(review_table.rows[2].cells[1],
    '团队负责人所在单位意见：\n\n'
    '□ 同意    □ 不同意    该团队的活动内容。\n\n\n'
    '签字：________________\n\n'
    '盖章：________________\n\n'
    '    年    月    日',
    False, Pt(11))

# ========== 保存文件 ==========
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           '附件3：2026年暑期社会实践活动院级项目申报书（已填写）.docx')
doc.save(output_path)
print(f'文件已保存：{output_path}')
print(f'共 {len(doc.paragraphs)} 个段落，{len(doc.tables)} 个表格')
