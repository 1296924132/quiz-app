#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
根据附件4 参赛作品申报书.docx模板格式，以反诈语音实时识别（三稿）为内容，
生成新的参赛作品申报书（已填写）.docx
完全按模板的表格结构和排版生成。
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ========== 页面设置（A4） ==========
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# ========== 辅助函数 ==========
def set_run_font(run, name='宋体', size=Pt(12), bold=False):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)

def set_para(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing_line=Pt(28),
             space_before=0, space_after=0, first_line_indent=None):
    p.alignment = alignment
    p.paragraph_format.line_spacing = spacing_line
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)

def add_run_para(text, font_name='宋体', size=Pt(12), bold=False,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing_line=Pt(28),
                 space_before=0, space_after=0, first_line_indent=None):
    p = doc.add_paragraph()
    set_para(p, alignment, spacing_line, space_before, space_after, first_line_indent)
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return p

def add_empty_para(spacing=28, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(spacing)

def set_cell_text(cell, text, name='宋体', size=Pt(10.5), bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=Pt(18)):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_run_font(run, name, size, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_table_borders(table, color="000000", sz="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        return
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

def set_grid_col_widths(table, widths_cm):
    """设置表格gridCol宽度（确保列宽在Word中正确显示）"""
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        return
    gridCols = tblGrid.findall(qn('w:gridCol'))
    for i, gc in enumerate(gridCols):
        if i < len(widths_cm):
            w_dxa = int(widths_cm[i] / 2.54 * 1440)
            gc.set(qn('w:w'), str(w_dxa))
    # 同时设置每行的列宽
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_cm):
                w_dxa = int(widths_cm[ci] / 2.54 * 1440)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{w_dxa}" w:type="dxa"/>')
                    tcPr.insert(0, tcW)
                else:
                    tcW.set(qn('w:w'), str(w_dxa))
                    tcW.set(qn('w:type'), 'dxa')

# ========== 封面部分 ==========

# 附件4 + 编号（右对齐）
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.line_spacing = Pt(24)
r = p.add_run('附件4')
set_run_font(r, '宋体', Pt(12))
r2 = p.add_run('\n编号：')
set_run_font(r2, '宋体', Pt(12))

add_empty_para(28, 2)

# 大标题
add_run_para('2026年重庆市大学生创新方法大赛', '黑体', Pt(22), True,
             WD_ALIGN_PARAGRAPH.CENTER, Pt(36), 6, 6)
add_run_para('作品申报书', '黑体', Pt(22), True,
             WD_ALIGN_PARAGRAPH.CENTER, Pt(36), 6, 6)

add_empty_para(28, 1)

# 类别行
add_run_para('（☑发明制作类、□工艺改进类、□创新设计类、□生活创意类）', '宋体', Pt(12), False,
             WD_ALIGN_PARAGRAPH.CENTER, Pt(28), 6, 12)

add_empty_para(28, 1)

# ---- 基本信息表格（无边框） ----
info_table = doc.add_table(rows=8, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(info_table)

info_data = [
    ('作品名称：', 'AI反诈通话语音实时识别与预警平台'),
    ('单    位：', '重庆警察学院'),
    ('申报团队：', '赵俊铭  姜沅其  钱丁豪  牟海涛  邓镰键'),
    ('创新方法导师：', '（请填写）'),
    ('专业导师：', '李小可  刘彦飞'),
    ('负 责 人：', '赵俊铭'),
    ('联系电话：', '（请填写）'),
    ('E-mail  ：', '（请填写）'),
]

for ri, (label, value) in enumerate(info_data):
    row = info_table.rows[ri]
    set_cell_text(row.cells[0], label, '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, Pt(26))
    set_cell_text(row.cells[1], value, '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, Pt(26))
set_grid_col_widths(info_table, [3.5, 11.14])

add_empty_para(28, 1)

# ---- 指导教师推荐理由（表格） ----
add_run_para('指导教师推荐理由：（希望评委关注的亮点，言简意赅、切忌浮夸、否则会影响成绩、80字内）',
             '宋体', Pt(10.5), False, WD_ALIGN_PARAGRAPH.LEFT, Pt(24), 6, 6)

reason_table = doc.add_table(rows=6, cols=1)
reason_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(reason_table)

reason_data = [
    '1、项目以AI反诈通话语音实时识别为核心，融合大模型表征与无监督异常检测，技术路线前沿，具有较强的创新性。',
    '2、项目构建的向量级隐私脱敏体系，兼顾检测精度与数据合规，解决了公安场景下敏感数据处理的现实难题。',
    '3、三级风险分级预警与闭环迭代机制设计完整，具备较高的实战应用价值和推广前景。',
    '4、团队成员专业基础扎实，分工明确，具备完成该项目研究的能力。',
    '5、该系统针对电信网络诈骗防范这一社会痛点，具有较强的社会效益和推广应用前景。',
    '',
]

for i, line in enumerate(reason_data):
    set_cell_text(reason_table.rows[i].cells[0], line, '宋体', Pt(10), False, WD_ALIGN_PARAGRAPH.LEFT, Pt(18))

add_empty_para(28, 2)

# 落款
add_run_para('重庆市大学生创新方法大赛组委会制', '宋体', Pt(12), False,
             WD_ALIGN_PARAGRAPH.CENTER, Pt(28), 6, 6)
add_run_para('2026年  月  日填写', '宋体', Pt(12), False,
             WD_ALIGN_PARAGRAPH.CENTER, Pt(28), 6, 6)

# ========== 说明部分 ==========
doc.add_page_break()

add_run_para('说      明', '黑体', Pt(14), True,
             WD_ALIGN_PARAGRAPH.CENTER, Pt(28), 12, 12)

instructions = [
    '1．申报者应在认真阅读此说明各项内容后按要求详细填写。',
    '2．申报者在填写申报作品时需详细填写表A、B。',
    '3．作品类别处只保留一项符合作品的类别或在方框内打√。',
    '4．编号由届次+学校编号+作品编号组成，组委会指定学校编号，参赛高校联络员统一填写作品编号。',
    '5．所有参赛作品必须按规定时间由高校联络员统一分类报送。',
    '6．所有参赛作品填写相关内容均需报送佐证材料扫描件，如已成立公司需提供企业营业执照；已申请专利需提供专利受理通知书、授权通知书及缴费凭证等相关佐证，发明人要有参赛学生名字。',
]

for inst in instructions:
    add_run_para(inst, '宋体', Pt(10.5), False,
                 WD_ALIGN_PARAGRAPH.LEFT, Pt(22), 4, 4, first_line_indent=0.5)

add_empty_para(28, 6)

# ========== A．参赛作品简介 ==========

add_run_para('A．参赛作品简介', '黑体', Pt(14), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(28), 12, 12)

add_run_para('作品名称：AI反诈通话语音实时识别与预警平台', '宋体', Pt(12), False,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(24), 6, 6)

add_run_para('作品简介（150字以内）：', '宋体', Pt(10.5), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(24), 6, 6)

brief = ('本项目构建实时合规精准的AI反诈通话语音实时识别与预警平台，'
         '通过语音标准化预处理、Whisper Encoder双维度特征提取、向量级脱敏处理，'
         '依托无监督自编码器实现异常检测，结合三级风险分级预警与闭环迭代机制，'
         '精准识别涉诈语音，对接公安等平台实现事中干预，兼顾检测效能与数据合规。')
add_run_para(brief, '仿宋', Pt(10.5), False,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(22), 2, 6, first_line_indent=0.74)

add_empty_para(28, 1)

add_run_para('应用的TRIZ理论（100字以内）：', '宋体', Pt(10.5), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(24), 6, 6)

triz_text = ('应用TRIZ分割原理将系统划分为语音输入、特征提取、数据脱敏、异常检测、'
            '预警输出五大功能模块；利用预先作用原理实现语音标准化预处理；'
            '借助中介物原理引入Whisper大模型与自编码器作为分析工具；'
            '采用参数变化原理以高斯噪声扰动实现数据脱敏。')
add_run_para(triz_text, '仿宋', Pt(10.5), False,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(22), 2, 6, first_line_indent=0.74)

add_empty_para(28, 2)

# ---- 专利申请表 ----
# 使用2列表格，纵向合并4行
pt = doc.add_table(rows=4, cols=2)
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(pt)

# 纵向合并第0列
pt.cell(0, 0).merge(pt.cell(3, 0))
pt.rows[0].cells[0].width = Cm(2.5)
pt.rows[0].cells[1].width = Cm(12.14)
set_cell_text(pt.rows[0].cells[0], '专利申请\n情况', '宋体', Pt(10), True, WD_ALIGN_PARAGRAPH.CENTER, Pt(18))
set_cell_text(pt.rows[0].cells[1], '□提出专利申请       申请号：________________', '宋体', Pt(10))
set_cell_text(pt.rows[1].cells[1], '申请日期______年______月______日', '宋体', Pt(10))
set_cell_text(pt.rows[2].cells[1], '□已获专利权批准     批准号：________________', '宋体', Pt(10))
set_cell_text(pt.rows[3].cells[1], '☑未提出专利申请', '宋体', Pt(10))
set_grid_col_widths(pt, [2.5, 12.14])

add_empty_para(28, 0)

# ---- 获奖表 ----
# 3列表格：第0列纵向合并，第1-2列每行横向合并
at = doc.add_table(rows=6, cols=3)
at.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(at)

# 合并第0列
at.cell(0, 0).merge(at.cell(5, 0))
at.rows[0].cells[0].width = Cm(2.5)
set_cell_text(at.rows[0].cells[0], '此作品参\n加本赛事\n之前获奖\n情况', '宋体', Pt(10), True, WD_ALIGN_PARAGRAPH.CENTER, Pt(18))

# 每行合并1-2列
for ri in range(6):
    at.cell(ri, 1).merge(at.cell(ri, 2))
    at.rows[ri].cells[1].width = Cm(12.14)

award_rows = [
    '是否获过奖：☑否    □是',
    '',
    '',
    '获过奖赛事名称：________________',
    '获奖级别：______________________',
    '',
]
for ri, txt in enumerate(award_rows):
    set_cell_text(at.rows[ri].cells[1], txt, '宋体', Pt(10))
set_grid_col_widths(at, [2.5, 12.14, 0])

add_empty_para(28, 0)

# ---- 成果转化表 ----
tt = doc.add_table(rows=3, cols=2)
tt.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(tt)

tt.cell(0, 0).merge(tt.cell(2, 0))
tt.rows[0].cells[0].width = Cm(2.5)
tt.rows[0].cells[1].width = Cm(12.14)
set_cell_text(tt.rows[0].cells[0], '成果转化\n情况', '宋体', Pt(10), True, WD_ALIGN_PARAGRAPH.CENTER, Pt(18))
set_cell_text(tt.rows[0].cells[1], '是否与企业对接：是 □   否 ☑', '宋体', Pt(10))
set_cell_text(tt.rows[1].cells[1], '对接企业名称：________________', '宋体', Pt(10))
set_cell_text(tt.rows[2].cells[1], '对接日期______年______月______日', '宋体', Pt(10))
set_grid_col_widths(tt, [2.5, 12.14])

add_empty_para(28, 2)

# ========== B．申报作品TRIZ理论应用情况 ==========
doc.add_page_break()

add_run_para('B．申报作品TRIZ理论应用情况', '黑体', Pt(14), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(28), 12, 12)

# ---- 说明框 ----
tip_table = doc.add_table(rows=8, cols=1)
tip_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(tip_table)

tip_items = [
    '本部分需要展现如下内容，请"亲们"仔细阅读、认真领会！',
    '1、运用TRIZ解决问题时的步骤；',
    '2、体现运用TRIZ各种创新工具解决问题时的自然诚恳态度；',
    '3、重要的不单是获得结果，而是运用TRIZ的分析和思考过程；',
    '4、如实记录方案产生的思维过程以及附带产生的想法；',
    '5、发明问题可以有很多解决方案，因此展现你的评价技能和选择最佳方案也很重要；',
    '6、完美展现、诠释你的最佳方案。',
    '学会运用创新方法分析、思考、解决问题远比获奖更重要！祝你取得好成绩！',
]
for i, item in enumerate(tip_items):
    bold = (i == 0 or i == 7)
    set_cell_text(tip_table.rows[i].cells[0], item, '宋体', Pt(10), bold, WD_ALIGN_PARAGRAPH.LEFT, Pt(18))

add_empty_para(28, 1)
add_run_para('以下步骤模板及案例仅供参考', '宋体', Pt(10.5), False,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(24), 6, 12)

# ===== 第一部分：问题描述 =====
add_run_para('第一部分：问题描述', '黑体', Pt(12), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(28), 12, 12)

add_run_para('1、项目概述', '宋体', Pt(10.5), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(24), 6, 6)

overview = (
    '（1）项目来源：电信网络诈骗已成为当前危害社会治安的突出问题，'
    'AI合成语音诈骗、冒充熟人诈骗等新型手段层出不穷，传统反诈手段难以有效应对。'
    '本项目源于公安实战需求，旨在解决反诈工作中"发现难、识别慢、拦截被动"的痛点。\n'
    '（2）问题描述：当前反诈工作面临三重难题——AI合成语音难以识别、'
    '新型诈骗话术迭代快导致规则库滞后、传统拦截规则误报率高影响用户体验。\n'
    '（3）技术方案：构建基于Whisper Encoder双维度特征提取与无监督自编码器异常检测的'
    'AI反诈通话语音实时识别与预警平台，实现"识别—研判—预警—总结"全链条闭环技术体系。\n'
    '（4）技术参数：语音输入统一为16kHz采样率、16bit量化、单声道WAV格式，'
    'Whisper Encoder提取高维embedding向量，自编码器以MSE损失函数计算重构误差，'
    '基于置信区间设定自适应检测阈值。'
)
add_run_para(overview, '仿宋', Pt(10.5), False,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(22), 2, 6, first_line_indent=0.74)

add_run_para('2、发明问题初始形势分析', '宋体', Pt(10.5), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(24), 12, 6)

init_analysis = (
    '系统工作原理：用户通话语音经标准化预处理后，由Whisper Encoder同步提取声学与语义双维度特征，'
    '生成embedding向量；经高斯噪声扰动与敏感维度掩码完成向量级脱敏；由无监督自编码器计算重构误差，'
    '基于误差大小进行三级风险分级预警。\n'
    '存在主要问题：①诈骗话术更新迅速，传统规则匹配方式滞后；'
    '②AI合成语音与真人语音高度相似，传统声纹检测难以区分；'
    '③通话语音包含大量个人敏感信息，数据处理需符合《个人信息保护法》及公安数据安全规范。\n'
    '限制条件：需在实时通话场景下完成全链路处理，端到端延迟需控制在可接受范围内；'
    '模型需在公安内网环境部署，资源受限。\n'
    '目前解决方案及不足：现有方案多为关键词匹配或基于有监督学习的分类模型，'
    '存在标注样本稀缺、对新型诈骗识别率低、误报率高等问题。'
)
add_run_para(init_analysis, '仿宋', Pt(10.5), False,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(22), 2, 6, first_line_indent=0.74)

# ===== 第二部分：系统分析 =====
add_run_para('第二部分：系统分析', '黑体', Pt(12), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(28), 12, 12)

add_run_para('3、系统分析', '宋体', Pt(10.5), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(24), 6, 6)

sys_analysis = (
    '（1）因果分析：诈骗通话的根本原因可追溯为——犯罪分子利用信息不对称、'
    '技术手段模仿可信身份、利用受害者心理弱点实施诈骗。'
    '本项目从技术层面切断诈骗实施链路，在通话过程中实时识别并预警。\n'
    '（2）九屏分析（系统/时间维度）：\n'
    '  超系统：传统电话诈骗→AI合成语音诈骗→深度伪造交互诈骗\n'
    '  系统：实时通话语音→Whisper+AE检测预警→多模态融合反诈\n'
    '  子系统：关键词过滤→深度学习模型→自适应认知智能\n'
    '（3）资源分析：充分利用语音资源（声学特征、语义特征）、'
    '计算资源（Whisper预训练模型、自编码器）、数据资源（正常通话样本）。\n'
    '（4）功能分析：语音输入层执行格式标准化与前端预处理；'
    '特征提取层完成声学与语义双模态同步提取；'
    '脱敏层实现敏感信息屏蔽；检测层完成异常识别；预警层输出风险等级。'
)
add_run_para(sys_analysis, '仿宋', Pt(10.5), False,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(22), 2, 6, first_line_indent=0.74)

# ===== 第三部分：运用TRIZ工具解决问题 =====
add_run_para('第三部分：运用TRIZ工具解决问题', '黑体', Pt(12), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(28), 12, 12)

add_run_para('4、TRIZ工具', '宋体', Pt(10.5), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(24), 6, 6)

triz_tools = (
    '（1）最终理想解（IFR）：理想状态下，系统应在完全不影响正常通话体验的前提下，'
    '实时、零误报地识别并拦截所有涉诈通话，且不存储任何个人隐私信息。\n'
    '（2）技术矛盾与发明原理：\n'
    '  矛盾1：检测精度与处理速度的矛盾（提高检测精度需要更多计算资源→降低处理速度）\n'
    '    解决：采用分割原理——将系统划分为独立的功能模块，并行处理；'
    '采用预先作用原理——语音标准化预处理，减少在线计算负担。\n'
    '  矛盾2：数据脱敏与保持语义完整性的矛盾（脱敏程度越高→可能丢失检测关键特征）\n'
    '    解决：采用参数变化原理——通过可控高斯噪声扰动而非直接删除敏感信息，'
    '在保护隐私的同时保留检测所需特征。\n'
    '（3）物理矛盾：系统需要处理大量数据以提高检测准确率，同时需要轻量化以保障实时性。\n'
    '    解决：采用中介物原理——引入Whisper大模型作为高效特征提取中介，'
    '将复杂计算集中在Encoder部分，下游仅需轻量级自编码器即可完成检测。\n'
    '（4）物-场分析：\n'
    '  原系统：通话语音（物质）→关键词过滤（场）→判定结果\n'
    '  问题：场的作用不足，无法识别复杂诈骗话术\n'
    '  改进：通话语音（物质）→Whisper特征提取（场1）→embedding向量→'
    '自编码器异常检测（场2）→精准判定结果'
)
add_run_para(triz_tools, '仿宋', Pt(10.5), False,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(22), 2, 6, first_line_indent=0.74)

# ===== 第四部分：技术方案整理与评价 =====
add_run_para('第四部分：技术方案整理与评价', '黑体', Pt(12), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(28), 12, 12)

add_run_para('5、全部技术方案及评价', '宋体', Pt(10.5), True,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(24), 6, 6)

solutions_text = (
    '方案1：基于Whisper Encoder + 无监督自编码器（AE）的异常检测方案（最终方案）\n'
    '  描述：采用Whisper预训练大模型提取双维度特征，以无监督自编码器进行异常检测。'
    '仅需正常通话语音即可训练，可识别未知新型诈骗与AI合成语音。\n'
    '  评价：检测精度高、泛化能力强、无需标注涉诈样本、可识别未知诈骗类型。'
    '综合评分最高，确定为本项目的最终方案。\n\n'
    '方案2：基于关键词匹配的传统规则方案\n'
    '  描述：构建诈骗关键词库，对通话内容进行关键词匹配和规则过滤。\n'
    '  评价：实现简单、可解释性强，但无法识别变种话术和AI合成语音，误报率高。\n\n'
    '方案3：基于有监督深度学习的分类方案（CNN/RNN）\n'
    '  描述：收集标注涉诈语音样本，训练有监督深度学习分类器。\n'
    '  评价：对已知诈骗类型识别效果好，但高度依赖标注样本质量，'
    '对新型诈骗识别能力弱，模型更新迭代成本高。\n\n'
    '专利预案：\n'
    '  拟申请发明专利："一种基于无监督自编码器的反诈通话语音实时识别方法与系统"——'
    '涵盖双维度特征提取方法、向量级数据脱敏方法、基于重构误差的异常检测及三级分级预警方法。\n\n'
    '最终确定方案：方案1（Whisper Encoder + 无监督自编码器异常检测方案）'
)
add_run_para(solutions_text, '仿宋', Pt(10.5), False,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(22), 2, 6, first_line_indent=0.74)

add_empty_para(28, 2)

# 注释
add_run_para('注：\n'
             '1. 本次填写如能用图的请尽量用图表示；\n'
             '2. 尽可能地运用多种TRIZ工具解题，但不局限于括号中所列的工具；\n'
             '3. 解决方案应为多种，确定最终方案应为一种或两种皆可。',
             '宋体', Pt(10.5), False,
             WD_ALIGN_PARAGRAPH.LEFT, Pt(22), 6, 6)

# ========== 保存 ==========
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           '附件4-已填写.docx')
doc.save(output_path)
print(f'文件已保存：{output_path}')
print(f'共 {len(doc.paragraphs)} 个段落，{len(doc.tables)} 个表格')
