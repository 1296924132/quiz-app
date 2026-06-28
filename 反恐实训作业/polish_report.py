# -*- coding: utf-8 -*-
"""
润色实训期末作业文档 - 将内容改写为更高级的形式
保留原有格式结构，仅修改文字内容
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy
import os

# ============================================================
# 第一步：读取原始文档，获取格式信息
# ============================================================
src_path = r'C:\Users\12969\Desktop\实训期末作业.doc'
# 无法直接读取 .doc，因此基于提取的内容重建

# ============================================================
# 第二步：创建新文档
# ============================================================
doc = Document()

# 设置默认字体为宋体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页面为 A4
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ============================================================
# 辅助函数
# ============================================================
def set_run_font(run, font_name='宋体', size=None, bold=False, color=None):
    """设置 run 的字体属性"""
    run.font.name = font_name
    run.font.bold = bold
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:eastAsia'), font_name)

def add_paragraph_with_format(doc_or_table, text, font_name='宋体', size=Pt(12),
                               bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                               space_before=0, space_after=0, line_spacing=1.5,
                               first_line_indent=None, color=None):
    """添加带格式的段落"""
    p = doc_or_table.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = first_line_indent

    run = p.add_run(text)
    set_run_font(run, font_name, size, bold, color)
    return p

def add_table_cell_text(cell, text, font_name='宋体', size=Pt(12), bold=False,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                         first_line_indent=None):
    """设置单元格文本及格式"""
    # 清除默认段落
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)

    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return cell

def set_cell_vertical_alignment(cell, align="center"):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    tcPr.append(vAlign)

def set_cell_width(cell, width_cm):
    """设置单元格宽度"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="0" w:type="dxa"/>')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')

def merge_cells(table, row_start, col_start, row_end, col_end):
    """合并单元格"""
    cell = table.cell(row_start, col_start)
    cell_to = table.cell(row_end, col_end)
    cell.merge(cell_to)
    return cell

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" '
            f'w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

# ============================================================
# 第三步：重建文档结构（保留原格式）
# ============================================================

# ---- 标题行 ----
add_paragraph_with_format(
    doc, '重庆警察学院实训报告',
    font_name='黑体', size=Pt(22), bold=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=12, space_after=6
)

add_paragraph_with_format(
    doc, '反恐怖现场处置（实训）',
    font_name='黑体', size=Pt(16), bold=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=6, space_after=12
)

# ---- 基本信息表格 ----
info_table = doc.add_table(rows=3, cols=4)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 设置列宽
for row in info_table.rows:
    for i, cell in enumerate(row.cells):
        if i == 0:
            set_cell_width(cell, 1.5)
        elif i == 1:
            set_cell_width(cell, 4.0)
        elif i == 2:
            set_cell_width(cell, 1.5)
        elif i == 3:
            set_cell_width(cell, 4.0)

# 设置边框
border_cfg = {'val': 'single', 'sz': '4', 'color': '000000'}
for row in info_table.rows:
    for cell in row.cells:
        set_cell_border(cell, top=border_cfg, bottom=border_cfg, left=border_cfg, right=border_cfg)

# 填充基本信息
info_data = [
    ['姓  名', '姜沅其', '学  号', 'WA250208'],
    ['区  队', '警务指挥与战术25-2', '实训项目', '反恐怖现场处置（实训）'],
    ['实训时间', '2026年上半年第10-15周', '成  绩', '']
]
for i, row_data in enumerate(info_data):
    for j, text in enumerate(row_data):
        bold = j in [0, 2]  # 标签加粗
        add_table_cell_text(info_table.cell(i, j), text, '宋体', Pt(12), bold,
                           WD_ALIGN_PARAGRAPH.CENTER)

# ---- 实训项目一目标 ----
add_paragraph_with_format(
    doc, '\n实训项目',
    font_name='黑体', size=Pt(14), bold=True,
    space_before=6, space_after=3
)

# ============================================================
# 润色后的内容 - 实训项目一目标
# ============================================================
polished_objective_1 = (
    '实训项目一：通过派出所反侵袭战法专项实训，系统厘清基层公安驻地立体化安全防卫体系的底层建构逻辑，'
    '深度掌握人防、物防、技防三位一体安防机制的协同运行原理、日常维护规程及风险前置管控要点；'
    '具备对驻地突发性暴力侵袭事件的瞬时风险研判能力与分区空间封控能力，娴熟完成最小作战单元的人员编组分工与同步协同处置，'
    '牢固树立"主动防御、前置避险"的核心战术思维，恪守"科学自保、依法反击"的实战准则，'
    '构建规范、高效、贴合基层所队实战需求的驻地反恐处突思维体系与行动范式。'
)

polished_objective_2 = (
    '实训项目二：通过街面巡逻防控应对冷兵器暴力袭击专项实训，精准掌握三人最小作战单元全链条标准化协同处置流程，'
    '熟练贯彻"寻掩蔽、分距离、呼叫支援—约束凶器、击打薄弱关节、制服控人—依法上铐、全方位安全搜身"三段式闭环处置战术；'
    '强化突发暴恐警情的秒级响应处置素养，严格遵循法定梯次用武裁量标准，'
    '在实战处置中优先实施周边群众隔离疏散、有效规避次生衍生危害，统筹兼顾群众生命安全与警务人员自身防护，'
    '全面锻造街面突发暴力事件规范化、专业化、法治化的实战处置能力。'
)

# 创建实训目标表格
obj_table = doc.add_table(rows=2, cols=2)
obj_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_width(obj_table.cell(0, 0), 1.5)
set_cell_width(obj_table.cell(0, 1), 10.0)
set_cell_width(obj_table.cell(1, 0), 1.5)
set_cell_width(obj_table.cell(1, 1), 10.0)

for row in obj_table.rows:
    for cell in row.cells:
        set_cell_border(cell, top=border_cfg, bottom=border_cfg, left=border_cfg, right=border_cfg)

add_table_cell_text(obj_table.cell(0, 0), '实训目标\n与要求', '宋体', Pt(12), True,
                   WD_ALIGN_PARAGRAPH.CENTER)
add_table_cell_text(obj_table.cell(0, 1), polished_objective_1, '宋体', Pt(12), False,
                   WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0.74))
# 设置垂直对齐
set_cell_vertical_alignment(obj_table.cell(0, 0), "center")

add_table_cell_text(obj_table.cell(1, 0), '', '宋体', Pt(12), False,
                   WD_ALIGN_PARAGRAPH.CENTER)
add_table_cell_text(obj_table.cell(1, 1), polished_objective_2, '宋体', Pt(12), False,
                   WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0.74))
# 合并第二行第一列
merge_cells(obj_table, 1, 0, 1, 0)

# ============================================================
# 实训项目一：派出所反侵袭战法实训
# ============================================================
add_paragraph_with_format(
    doc, '\n实训项目一',
    font_name='黑体', size=Pt(14), bold=True,
    space_before=6, space_after=3
)

# 项目信息表格
proj1_table = doc.add_table(rows=4, cols=2)
proj1_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_width(proj1_table.cell(0, 0), 2.0)
set_cell_width(proj1_table.cell(0, 1), 9.5)
set_cell_width(proj1_table.cell(1, 0), 2.0)
set_cell_width(proj1_table.cell(1, 1), 9.5)
set_cell_width(proj1_table.cell(2, 0), 2.0)
set_cell_width(proj1_table.cell(2, 1), 9.5)
# 第4行合并
set_cell_width(proj1_table.cell(3, 0), 2.0)
set_cell_width(proj1_table.cell(3, 1), 9.5)

for row in proj1_table.rows:
    for cell in row.cells:
        set_cell_border(cell, top=border_cfg, bottom=border_cfg, left=border_cfg, right=border_cfg)

add_table_cell_text(proj1_table.cell(0, 0), '项目名称', '宋体', Pt(12), True,
                   WD_ALIGN_PARAGRAPH.CENTER)
set_cell_vertical_alignment(proj1_table.cell(0, 0), "center")
add_table_cell_text(proj1_table.cell(0, 1), '派出所反侵袭战法实训', '宋体', Pt(12), False,
                   WD_ALIGN_PARAGRAPH.LEFT)

add_table_cell_text(proj1_table.cell(1, 0), '时间地点', '宋体', Pt(12), True,
                   WD_ALIGN_PARAGRAPH.CENTER)
set_cell_vertical_alignment(proj1_table.cell(1, 0), "center")
add_table_cell_text(proj1_table.cell(1, 1), '2026年上半年第10-11周    勤勉楼教室、警训馆', '宋体', Pt(12), False,
                   WD_ALIGN_PARAGRAPH.LEFT)

add_table_cell_text(proj1_table.cell(2, 0), '主要内容', '宋体', Pt(12), True,
                   WD_ALIGN_PARAGRAPH.CENTER)
set_cell_vertical_alignment(proj1_table.cell(2, 0), "center")
add_table_cell_text(proj1_table.cell(2, 1), '结合驻地反侵袭标准化现场处置流程，推演某派出所反侵袭战法实训场景，'
                   '系统训练基层警务人员在驻地面临突发暴力侵袭时的全流程应对能力。',
                   '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0.74))

add_table_cell_text(proj1_table.cell(3, 0), '实训过程\n与内容记录', '宋体', Pt(12), True,
                   WD_ALIGN_PARAGRAPH.CENTER)
set_cell_vertical_alignment(proj1_table.cell(3, 0), "center")

# ---- 实训一过程记录（润色版） ----
training1_process = (
    '本次实训以"某派出所凌晨遭遇驾车冲撞并伴随持刀袭警"复合警情为预设背景，'
    '在警训馆内模拟派出所值班室、接待大厅及外部通道等重点区域，'
    '围绕威胁识别评估、先期预警响应、空间分区管控、最小作战单元反击、'
    '战术紧急医疗救援和战后系统复盘等环节，开展全流程、全要素实战化训练。'
    '实训始终坚持"安全第一、依法处置、协同配合、快速控制"的核心原则，'
    '着力提升派出所一线民警在突发暴力侵袭情境下的应急反应效能与团队协同作战能力。\n\n'
    '第一阶段：威胁识别与预警响应\n'
    '面对车辆异常冲撞、人员持械逼近、群众受惊聚集等多重复合突发情况，'
    '值班民警首先通过精准观察嫌疑人行为特征、科学判断危险源分布位置、'
    '系统研判侵袭发展方向等方式，迅速甄别警情性质与危险等级，'
    '并第一时间发出预警信号。现场人员按照既定岗位分工高效联动：'
    '接警人员即刻上报警情态势并请求增援，前端处置人员迅速完成防护准备，'
    '后方人员有序引导无关人员撤离至安全区域。'
    '通过本阶段训练，本人深刻认识到：派出所反侵袭处置断不可待危险完全显现后方被动作出反应，'
    '而应当在异常迹象初现端倪之际即行预判、预提醒、预布控，'
    '方能为后续处置环节赢得战略主动与宝贵先机。\n\n'
    '第二阶段：空间控制与坚壁清野\n'
    '演练中，参训人员依据派出所内部空间布局特征，对接待区、通道口、门厅等关键节点实施分区控制，'
    '巧妙利用桌椅、门禁系统、警戒带等现有条件划定相对安全区域，'
    '有效压缩嫌疑人活动空间，阻断其接近群众和民警的路线。'
    '同时，现场人员有序组织群众疏散撤离，严防人员围观、逆行、滞留所引发的二次风险。'
    '在此环节中，本人深切体悟到：空间控制绝非简单的"后退躲避"，'
    '而是通过科学站位、明确分工与通道动态管理，将混乱无序的现场转化为可观察、可研判、可处置的受控状态。\n\n'
    '第三阶段：最小作战单元反击战术\n'
    '面对嫌疑人持械冲击的严峻态势，三人最小作战单元严格遵循"观察警戒、正面牵制、侧翼控制、协同制服"的战术思路开展处置。'
    '各成员之间始终保持安全作战距离与相互支援角度：一人专职负责口头警告与吸引注意力，'
    '一人专注防护与物理阻隔，一人伺机寻找控制时机，并在条件成熟时协同完成对嫌疑人的有效控制。'
    '全过程强调依法规范使用警械，坚决杜绝单人冒进与盲目近身，切实做到处置有序、控制有效、风险可控。'
    '通过反复训练，本人进一步深化了对团队战术价值的理解：在高压力警情中，个人能力固然重要，'
    '但更为关键的是队形稳定、口令明确、动作协调与相互保护所形成的整体作战效能。\n\n'
    '第四阶段：战术医疗与战后复盘\n'
    '嫌疑人被成功控制后，现场人员继续保持高度警戒，对嫌疑人实施全面安全检查，'
    '确认其是否藏匿其他危险物品，同时对受伤民警与群众开展初步救护，'
    '并依法保护现场、固定保全证据、等待后续支援力量到达。'
    '复盘环节中，教官围绕预警是否及时、分工是否清晰、站位是否合理、控制是否规范、语言警告是否到位等维度进行系统点评。'
    '通过深入复盘，本人清醒认识到自身在复杂警情处置中尚存在观察范围不够全面、口令表达不够果断、'
    '与队友衔接配合不够紧密等不足之处。'
    '今后应进一步加强警情综合研判、团队协同配合以及现场心理稳定能力等方面的专项训练，'
    '持续提升面对突发暴力侵袭时的综合处置水平与实战应对能力。'
)
add_table_cell_text(proj1_table.cell(3, 1), training1_process, '宋体', Pt(12), False,
                   WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0.74))

# ============================================================
# 实训项目二：街面巡逻防控应对冷兵器袭击实训
# ============================================================
add_paragraph_with_format(
    doc, '\n实训项目二',
    font_name='黑体', size=Pt(14), bold=True,
    space_before=6, space_after=3
)

proj2_table = doc.add_table(rows=4, cols=2)
proj2_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_width(proj2_table.cell(0, 0), 2.0)
set_cell_width(proj2_table.cell(0, 1), 9.5)
set_cell_width(proj2_table.cell(1, 0), 2.0)
set_cell_width(proj2_table.cell(1, 1), 9.5)
set_cell_width(proj2_table.cell(2, 0), 2.0)
set_cell_width(proj2_table.cell(2, 1), 9.5)
set_cell_width(proj2_table.cell(3, 0), 2.0)
set_cell_width(proj2_table.cell(3, 1), 9.5)

for row in proj2_table.rows:
    for cell in row.cells:
        set_cell_border(cell, top=border_cfg, bottom=border_cfg, left=border_cfg, right=border_cfg)

add_table_cell_text(proj2_table.cell(0, 0), '项目名称', '宋体', Pt(12), True,
                   WD_ALIGN_PARAGRAPH.CENTER)
set_cell_vertical_alignment(proj2_table.cell(0, 0), "center")
add_table_cell_text(proj2_table.cell(0, 1), '街面巡逻防控应对冷兵器袭击实训', '宋体', Pt(12), False,
                   WD_ALIGN_PARAGRAPH.LEFT)

add_table_cell_text(proj2_table.cell(1, 0), '时间地点', '宋体', Pt(12), True,
                   WD_ALIGN_PARAGRAPH.CENTER)
set_cell_vertical_alignment(proj2_table.cell(1, 0), "center")
add_table_cell_text(proj2_table.cell(1, 1), '2026年上半年第14-15周    勤勉楼教室、警训馆', '宋体', Pt(12), False,
                   WD_ALIGN_PARAGRAPH.LEFT)

add_table_cell_text(proj2_table.cell(2, 0), '主要内容', '宋体', Pt(12), True,
                   WD_ALIGN_PARAGRAPH.CENTER)
set_cell_vertical_alignment(proj2_table.cell(2, 0), "center")
add_table_cell_text(proj2_table.cell(2, 1), '【案例推演】以2024年某市商业街持刀无差别砍人事件为原型'
                   '（嫌疑人持50厘米户外刀具致3人受伤），三人巡逻单元（盾手甲、棍手乙、枪手丙组长）'
                   '以最小作战单元编组开展案例推演与实战化训练。',
                   '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0.74))

add_table_cell_text(proj2_table.cell(3, 0), '实训过程\n与内容记录', '宋体', Pt(12), True,
                   WD_ALIGN_PARAGRAPH.CENTER)
set_cell_vertical_alignment(proj2_table.cell(3, 0), "center")

# ---- 实训二过程记录（润色版） ----
training2_process = (
    '本次实训以商业街人员密集区域发生持刀无差别伤人重大警情为背景场景，'
    '全程模拟三人巡逻单元在街面巡逻执勤过程中发现突发暴力犯罪后，'
    '如何科学开展快速反应、现场封控、群众疏散、协同处置、伤员救护及后续移交等关键环节。'
    '该类警情具有突发性极强、危害辐射面广、群众恐慌情绪显著、现场干扰因素错综复杂等突出特点，'
    '若处置稍有迟滞，极易酿成更为严重的人员伤亡与社会负面影响。'
    '因此，本次训练将重点聚焦于巡逻警力在复杂公共空间中的应急态势研判能力、'
    '最小作战单元协同作战效能以及现场秩序整体掌控能力的全面锤炼。\n\n'
    '实训开始前，教官结合典型案例对冷兵器袭击警情的规律特征进行了系统深入分析。'
    '持刀伤人事件往往具有突发性特征，嫌疑人攻击目标具有高度不确定性，移动路径难以预先判断，'
    '现场群众极易因恐慌情绪引发奔跑、摔倒、围观滞留、不当拍摄等次生干扰行为。'
    '面对此类警情态势，巡逻民警既要迅速接近事发核心区域、果断制止正在实施的不法侵害，'
    '又绝不能因盲目冲动而导致自身负伤或进一步扩大风险敞口。'
    '随后，参训人员以三人巡逻小组为基本作战单元，严格遵循"发现警情、判断危险、疏散群众、'
    '控制现场、协同处置、救护移交"的标准流程展开系统化演练。\n\n'
    '第一阶段：警情发现与快速研判\n'
    '演练中，巡逻单元在商业街模拟区域发现群众异常惊慌奔跑态势，并清晰听到有人呼救。'
    '嫌疑人手持长刀在街面挥舞，已造成群众受伤，且存在继续攻击他人的高度可能性。'
    '三名队员迅速判明该警情属于正在发生的严重暴力犯罪，必须立即采取果断措施制止不法侵害。'
    '小组长第一时间下达处置指令，明确一名队员负责上报警情态势、请求增援力量并联系医疗急救，'
    '一名队员负责有序疏散群众、开辟安全撤离通道，一名队员负责保持警戒状态、持续观察嫌疑人动态。'
    '通过本环节训练，本人深刻认识到：街面巡逻防控的首要核心能力在于"精准发现问题"与"科学判断风险"。'
    '唯有在最短时限内精准认清警情性质，方能从根本上避免处置迟缓与现场失序的被动局面。\n\n'
    '第二阶段：群众疏散与危险区域控制\n'
    '商业街作为典型人员密集场所，现场群众基数庞大、移动方向纷繁复杂，'
    '若不及时实施有效疏散，嫌疑人可能继续伤害无辜群众，甚至引发踩踏等连锁次生灾害。'
    '演练中，巡逻人员通过简短明确的口令引导群众向背离嫌疑人方向有序撤离，'
    '并反复提醒周边人员不得围观、不得拍摄、不得靠近危险警戒区域。'
    '同时，队员充分利用现场摊位、墙体结构、隔离护栏等环境条件，'
    '初步划定危险区、缓冲区与安全区的三级防控体系，最大限度将嫌疑人与群众实施物理隔离。'
    '该阶段的训练使本人深刻领悟到：街面警情处置的根本核心目标是保护人民群众生命安全。'
    '制服嫌疑人固然是重要任务，但在嫌疑人尚未被完全控制之前，'
    '最大限度降低群众暴露于危险范围之内的风险，同样具有同等重要的战略意义。\n\n'
    '第三阶段：三人最小作战单元协同逼控\n'
    '面对持刀嫌疑人的严重威胁，三名队员摒弃单人强行近身的冒险做法，'
    '而是通过科学队形配合逐步压缩嫌疑人活动空间。正面队员保持安全警戒距离，'
    '持续发出严正警告，要求嫌疑人立即停止违法犯罪行为并放下凶器；'
    '侧翼队员根据现场态势动态调整控制角度，防止嫌疑人突然转向攻击群众或突破防线；'
    '后方队员专职负责观察周边环境变化并随时准备支援队友，严防视野盲区与新的危险点出现。'
    '三名队员在行动全程始终保持相互可见、相互支援的战术状态，'
    '既形成对嫌疑人的强大心理震慑效应，又充分保障自身安全。'
    '通过系统训练，本人深刻体会到：最小作战单元的关键要义不仅在于"三个人在一起"，'
    '更在于三个人之间能够形成明确分工、有效掩护与同步行动的有机整体。\n\n'
    '第四阶段：依法使用警械与控制嫌疑人\n'
    '演练中，当嫌疑人不听警告并继续持刀逼近时，处置人员根据现场危险程度与法律规定，'
    '依法采取必要的警械控制措施。队员之间通过统一口令精准把握处置时机，'
    '在确保群众已安全撤离、队友站位相对安全的前提条件下，协同上前完成有效控制。'
    '控制完成后，处置人员立即对嫌疑人实施约束与全面安全检查，'
    '确认其是否仍藏匿其他刀具或危险物品，并保持对其持续严密看管。'
    '该环节的训练使本人深刻认识到：依法用警绝非消极等待的托词，'
    '而是在法律框架内果断、规范、有效地制止违法犯罪行为的行动准则。'
    '面对正在发生的严重暴力侵害，处置人员必须具备敢于担当的政治勇气与果断行动的专业素养，'
    '同时亦须高度重视程序规范与证据保全意识。\n\n'
    '第五阶段：伤员救助、现场保护与后续移交\n'
    '嫌疑人被有效控制后，巡逻单元立即转入后续处置程序。一名队员继续警戒嫌疑人，'
    '严防其反抗或再次实施攻击；一名队员迅速查看受伤群众伤情，开展初步止血救护、'
    '安抚情绪并耐心等待医疗救援力量抵达；另一名队员维护现场整体秩序，'
    '严禁无关人员进入核心处置区域，并向后续增援力量全面说明现场经过、嫌疑人控制情况、'
    '伤员分布位置及遗留物品状况。通过本环节训练，本人深刻认识到：'
    '街面突发暴力警情的处置是一项高度连续性的系统工程，从发现警情到控制嫌疑人，'
    '再到救助伤员、保护现场与移交后续力量，每一个环节都直接关系到整体处置质量与最终执法效果。\n\n'
    '在复盘点评环节，教官重点指出了处置过程中需着力改进的几个关键问题：'
    '一是警情判断必须果断坚决，绝不能因现场混乱而犹豫不决；'
    '二是群众疏散必须置于优先地位，绝不能使群众继续滞留于嫌疑人攻击范围之内；'
    '三是队员站位必须科学合理，杜绝出现队形过散、相互遮挡或支援不及时等战术缺陷；'
    '四是口令必须清晰有力，既要对嫌疑人形成有效震慑，也要引导群众有序避险；'
    '五是控制嫌疑人后仍须保持高度警戒，严防因忽视二次危险而酿成新的被动局面。'
    '通过系统复盘，本人清醒认识到自身在处置过程中尚存在对现场整体环境观察不够细致、'
    '对群众疏散方向判断不够迅速、与队友之间眼神交流和口令配合不够默契等方面的现实不足。'
    '这些深刻教训时刻提醒本人：实战能力的真正提升，必须建立在反复训练与认真总结的坚实基础之上。\n\n'
    '通过本次街面巡逻防控应对冷兵器袭击实训，本人对巡逻防控工作的实战价值有了更加深刻的理解与认知。'
    '街面巡逻不仅是维护社会面治安秩序的重要基础手段，更是公安机关快速发现、快速反应、'
    '快速处置突发警情的关键力量支撑。面对冷兵器袭击等严重暴力犯罪，'
    '巡逻民警必须具备高度警觉的防范意识、稳定的心理素质与规范的处置能力。'
    '此次实训使本人深刻认识到：作为未来的警务工作者，既要有冲锋在前的无畏勇气，'
    '也要有科学处置的专业智慧；既要敢于同违法犯罪行为作坚决斗争，'
    '也要善于保护群众、保护队友、保护自己。'
    '今后本人将持续加强队形协同、警械使用、语言控制、现场观察与心理承压等专项训练，'
    '不断提升自身综合实战能力与专业素养。'
)
add_table_cell_text(proj2_table.cell(3, 1), training2_process, '宋体', Pt(12), False,
                   WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0.74))

# ============================================================
# 实训总结（润色版）
# ============================================================
add_paragraph_with_format(
    doc, '\n实训总结',
    font_name='黑体', size=Pt(14), bold=True,
    space_before=6, space_after=3
)

conclusion = (
    '通过上述两个实训项目的系统学习与实战演练，本人更加深刻地认识到：'
    '公安实战训练的根本意义远不止于完成既定的规定动作，而是在于高度接近真实警情的复杂环境中，'
    '全面培养临场态势判断、团队协同配合、依法规范处置与风险动态管控等综合实战能力。'
    '派出所反侵袭战法实训重点聚焦于基层警务场所在遭遇突发暴力侵袭时的有效防护与果断反制，'
    '街面巡逻防控应对冷兵器袭击实训则侧重于公共空间中突发暴力犯罪的及时发现、有效控制与科学处置。'
    '两个实训项目虽然应用场景各有不同，但均深刻体现了公安工作"人民至上、生命至上、依法规范、协同作战"的根本原则与核心要求。\n\n'
    '通过系统实训，本人深刻认识到：警务处置从来不是单纯依靠个人胆识和勇气所能完成的简单任务，'
    '而是必须建立在严格规范流程、深厚战术素养、牢固法治意识与默契团队配合基础之上的高度综合性行动。'
    '面对危险警情，警务人员既不可畏缩迟疑、贻误战机，亦不可盲目冒进、置安全于不顾；'
    '既要坚决果断地制止违法犯罪行为，也要最大限度降低群众、队友及自身受到伤害的可能性。'
    '此次实训使本人更加明确了未来学习训练的努力方向：\n\n'
    '一是着力提高警情综合研判能力，切实做到早发现、早判断、早处置；\n'
    '二是持续加强最小作战单元协同训练，不断提升与队友之间的战术默契配合程度；\n'
    '三是强化依法用警规范意识，确保处置有法律依据、行为有明确边界、过程有严格规范；\n'
    '四是着力提升心理承压与情绪调控能力，在复杂紧张环境下始终保持冷静、果断与清醒的战术思维。\n\n'
    '作为警务指挥与战术专业的学员，本人深知本次期末实训不仅是一次课程考核的形式环节，'
    '更是对未来公安职业实战能力的一次全面检验与提前预演。'
    '今后，本人将以更高标准严格要求自己，将课堂理论、警务技能与实战训练有机结合，'
    '在每一次训练中认真查找不足、补齐能力短板、锤炼过硬作风，'
    '努力成长为政治立场坚定、工作作风过硬、业务技能精湛、敢于担当尽责的新时代高素质警务专门人才。'
)

add_paragraph_with_format(
    doc, conclusion,
    font_name='宋体', size=Pt(12), bold=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    space_before=3, space_after=3,
    first_line_indent=Cm(0.74)
)

# ---- 签名区 ----
add_paragraph_with_format(doc, '', '宋体', Pt(12), space_before=6, space_after=6)
add_paragraph_with_format(
    doc, '教师签字：',
    font_name='宋体', size=Pt(12), bold=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    space_before=12, space_after=3
)
add_paragraph_with_format(
    doc, '                                      年    月    日',
    font_name='宋体', size=Pt(12), bold=False,
    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    space_before=3, space_after=3
)

# ============================================================
# 保存文档
# ============================================================
output_path = r'C:\Users\12969\Desktop\实训期末作业（润色版）.docx'
doc.save(output_path)
print(f'文档已保存至：{output_path}')
print('润色完成！')
