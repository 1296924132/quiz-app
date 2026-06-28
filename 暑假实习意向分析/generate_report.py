import json, os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT_DIR = r"D:\TRAE文件\暑假实习意向分析"
json_path = os.path.join(OUTPUT_DIR, 'summary_data.json')

with open(json_path, 'r', encoding='utf-8') as f:
    data = json.load(f)

doc = Document()

# ---- Style defaults ----
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== Title Page =====
for _ in range(4):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('2026年暑假实习意向调查\n分析报告')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

doc.add_paragraph('')
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'统计日期：2026年6月')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(6):
    doc.add_paragraph('')

footer_note = doc.add_paragraph()
footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_note.add_run('涵盖交管、侦查、刑技、战术、政工、治安、涉外、网安等8个专业12个区队')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ===== Helper Functions =====
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    return p

def set_cell(cell, text, bold=False, color=None, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 表格单元格边距
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), '40')
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def shade_cells(row, color='D6E4F0'):
    for cell in row.cells:
        shading = cell._tc.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): color
        })
        shading.append(shading_elm)

def add_table(headers, rows_data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        set_cell(hdr.cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
    shade_cells(hdr, '2F5496')

    # Data
    for row_data in rows_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            set_cell(row.cells[i], val, size=10)

    # Alternate row shading
    for idx, row in enumerate(table.rows[1:], 1):
        if idx % 2 == 0:
            shade_cells(row, 'F2F2F2')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph('')
    return table

# ===== 1. 概述 =====
add_heading('一、调查概述', 1)
add_body(f'本次暑假实习意向调查覆盖全校 8 个专业的 12 个区队，共计 {data["total"]} 名学生。'
         f'调查内容主要为学生是否愿意参加暑期实习以及意向实习区域分布。')

add_body(f'调查结果显示：\n'
         f'  ● 有意向参加暑假实习的人数：{data["willing"]} 人\n'
         f'  ● 无意向参加暑假实习的人数：{data["unwilling"]} 人\n'
         f'  ● 总体意向率：{data["willing_rate"]}')

# ===== 2. 各区队意向统计 =====
add_heading('二、各区队意向分析', 1)
add_body(f'从各区队数据来看，意向率差异显著。'
         f'意向率最高的区队为 {data["top_team"]}，最低的为网安25-2（11.8%）。')

headers = ['排名', '区队', '总人数', '有意向', '无意向', '意向率']
rows = []
for i, s in enumerate(data['team_rank'], 1):
    ts = next(t for t in data['team_stats'] if t['区队'] == s)
    rows.append([i, ts['区队'], ts['总人数'], ts['有意向'], ts['无意向'], f'{ts["意向率"]:.1f}%'])
rows.append(['', '合计', data['total'], data['willing'], data['unwilling'], data['willing_rate']])

add_table(headers, rows, [1.5, 3, 2, 2, 2, 2])

# 分析文字
top3 = data['team_rank'][:3]
bottom3 = data['team_rank'][-3:]
top3_str = '、'.join([f'{t}' for t in top3])
bottom3_str = '、'.join([f'{t}' for t in bottom3])
add_body(f'排名前三的区队为：{top3_str}，意向率均超过 50%，展现出较强的实习意愿。'
         f'排名靠后的区队为：{bottom3_str}，意向率低于 20%，需要进一步了解原因并鼓励参与。')

# ===== 3. 整体意向率分析 =====
add_heading('三、整体意向率分析', 1)
w_rate = float(data['willing_rate'].replace('%', ''))
if w_rate > 50:
    analysis = f'整体意向率为 {data["willing_rate"]}，超过半数学生有意向参加暑假实习，反映出学生对实践锻炼的积极态度。'
elif w_rate > 30:
    analysis = f'整体意向率为 {data["willing_rate"]}，有一定比例的学生对暑假实习感兴趣，但仍有较大提升空间。'
else:
    analysis = f'整体意向率为 {data["willing_rate"]}，意向率偏低，建议加强实习宣传和动员工作。'

add_body(analysis)
add_body(f'在 {data["total"]} 名学生中，{data["willing"]} 人表示愿意参加暑假实习，'
         f'{data["unwilling"]} 人目前无意向参加。')
add_body(f'各区队中，侦查25-1以73.3%的意向率位居首位，远超平均水平；'
         f'网安25-2为11.8%，与首位相差61.5个百分点，区队间差距较大。')

# ===== 4. 意向地点分析 =====
add_heading('四、意向实习地点分析', 1)
add_body(f'在 {data["willing"]} 名有意向的学生中，实习意向覆盖全市多个区县。'
         f'最受欢迎的实习地点为 {data["top_location"]}（{data["top_location_cnt"]}人选择），'
         f'反映出学生对特定区域的偏好。')

loc_headers = ['排名', '意向地点', '人数', '占比']
loc_rows = []
for rank, (loc, cnt, rate) in enumerate(data['location_rank'], 1):
    loc_rows.append([rank, loc, cnt, rate])

add_table(loc_headers, loc_rows[:15], [1.5, 3.5, 2, 2])

# 分析
top_locs = data['location_rank'][:5]
top_loc_names = '、'.join([f'{loc}（{cnt}人）' for loc, cnt, _ in top_locs])
add_body(f'意向人数最多的前五个地点为：{top_loc_names}。'
         f'其中{data["top_location"]}为最热门选择。')

# 区域特点分析
urban = ['渝中区', '渝北区', '南岸区', '江北区', '九龙坡区', '沙坪坝区', '大渡口区', '巴南区', '北碚区']
suburban = ['万州区', '涪陵区', '永川区', '合川区', '江津区', '璧山区', '长寿区', '南川区', '綦江区', '大足区', '铜梁区', '荣昌区', '潼南区']
counties = ['云阳县', '奉节县', '丰都县', '忠县', '石柱县', '秀山县', '彭水县', '酉阳县', '巫溪县', '垫江县', '武隆区', '黔江区']

urban_total = sum(cnt for loc, cnt, _ in data['location_rank'] if loc in urban)
suburban_total = sum(cnt for loc, cnt, _ in data['location_rank'] if loc in suburban)
counties_total = sum(cnt for loc, cnt, _ in data['location_rank'] if loc in counties)
liangjiang = sum(cnt for loc, cnt, _ in data['location_rank'] if '两江' in loc)
gaoxin = sum(cnt for loc, cnt, _ in data['location_rank'] if '高新' in loc)
other = sum(cnt for loc, cnt, _ in data['location_rank'] if loc not in urban + suburban + counties and '两江' not in loc and '高新' not in loc)
urban_total += liangjiang + gaoxin + other

add_body(f'从区域类型看：\n'
         f'  ● 主城区意向：约 {urban_total} 人\n'
         f'  ● 区县意向：约 {suburban_total + counties_total} 人\n'
         f'  ● 其他/未明确：约 {other} 人')
add_body('总体来看，学生实习意向分布较广，主城区与区县均有覆盖，'
         '建议根据学生意向进行合理分配和安排。')

# ===== 5. 结论与建议 =====
add_heading('五、结论与建议', 1)

add_heading('5.1 主要结论', 2)
add_body(f'1. 全校 {data["total"]} 名学生中，{data["willing"]} 人（{data["willing_rate"]}）有意向参加暑假实习。')
add_body(f'2. 区队间差异较大：侦查25-1意向率最高（73.3%），网安25-2最低（11.8%）。')
add_body(f'3. 最受欢迎的实习地点为 {data["top_location"]}，主城区意向集中度较高。')
add_body(f'4. 共有 {len(data["location_rank"])} 个不同的实习意向地点，分布广泛。')

add_heading('5.2 工作建议', 2)
add_body('1. 加强宣传动员：针对意向率较低的区队（尤其是网安25-2、涉外25-1、政工25-1），'
         '建议专门开展实习动员会，强调实习对专业能力提升的帮助。')
add_body('2. 合理调配资源：根据意向地点分布，提前与相关单位对接，'
         '确保热门地区（如沙坪坝区、南岸区、两江新区）有充足的实习岗位。')
add_body('3. 关注个性化需求：对于填写"无要求"或"待定"的学生，安排专人跟进了解其顾虑，'
         '提供针对性的实习推荐。')
add_body('4. 做好统筹协调：区队间意向率差异较大，建议在学院层面统筹协调，'
         '将有意向的学生合理分配到各实习基地。')

# Footer
doc.add_paragraph('')
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— 报告结束 —')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

report_path = os.path.join(OUTPUT_DIR, '暑假实习意向分析报告.docx')
doc.save(report_path)
print(f'Word报告已保存: {report_path}')
